from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_table(doc, headers, rows, header_color="1F3864", alt_color="EBF3FB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = alt_color if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    return table

def add_info_box(doc, text, bg="D6E4F7", text_color=(31, 56, 100)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*text_color)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg)
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(220, 50, 50)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

# ════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ════════════════════════════════════════
# PAGE DE TITRE
# ════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("FICHE TECHNIQUE")
run.font.size = Pt(30)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 56, 100)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Déploiement Cloud — Système RAG")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(70, 130, 180)

doc.add_paragraph()
line = doc.add_paragraph("─" * 60)
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

badge_p = doc.add_paragraph()
badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
badge_run = badge_p.add_run("  3 OPTIONS — DU 100% GRATUIT AU PROFESSIONNEL  ")
badge_run.font.bold = True
badge_run.font.size = Pt(12)
badge_run.font.color.rgb = RGBColor(255, 255, 255)
pPr = badge_p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), '1F3864')
pPr.append(shd)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = info.add_run(f"Projet d'Entreprise — Chatbot RAG Documentaire\nDate : {datetime.date.today().strftime('%d %B %Y')}")
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ════════════════════════════════════════
# PROBLEMATIQUE
# ════════════════════════════════════════
add_heading(doc, "1. Problématique du Déploiement Cloud", level=1, color=(31, 56, 100))

doc.add_paragraph(
    "Le système RAG fonctionne en local avec Ollama (mistral:7b + nomic-embed-text). "
    "Pour un déploiement cloud, le défi principal est la mémoire RAM : "
    "Ollama + mistral:7b nécessite environ 4 à 8 Go de RAM, ce qui dépasse les ressources "
    "offertes par la majorité des free tiers cloud (512 Mo à 1 Go)."
)

doc.add_paragraph()
add_table(doc,
    ["Contrainte", "Impact"],
    [
        ("mistral:7b → ~4-8 Go RAM", "Élimine Render free tier (512 Mo), Railway free tier (512 Mo), Heroku (512 Mo)"),
        ("nomic-embed-text → ~274 Mo", "Modèle d'embedding additionnel à charger en mémoire"),
        ("ChromaDB persistant", "Nécessite un stockage disque partagé ou un volume persistant"),
        ("Temps de démarrage Ollama", "Cold start de ~30 sec — incompatible avec les services serverless"),
    ],
    header_color="8B1A1A",
    alt_color="FDECEA"
)

doc.add_paragraph()

# ════════════════════════════════════════
# VUE D'ENSEMBLE DES OPTIONS
# ════════════════════════════════════════
add_heading(doc, "2. Vue d'Ensemble des Options", level=1, color=(31, 56, 100))

add_table(doc,
    ["Option", "Architecture", "Coût", "Complexité", "Recommandé pour"],
    [
        ("A — Oracle Cloud", "Ollama complet sur VM gratuite", "0 €", "Moyenne", "Projet d'entreprise, confidentialité maximale"),
        ("B — Hybride Groq/Gemini", "App cloud + API LLM gratuite", "0 €", "Faible", "Déploiement rapide, scalabilité"),
        ("C — VPS Payant", "Ollama sur serveur dédié", "5-20 €/mois", "Faible", "Production, haute disponibilité"),
    ],
    header_color="1F3864",
    alt_color="EBF3FB"
)

doc.add_paragraph()

# ════════════════════════════════════════
# OPTION A
# ════════════════════════════════════════
add_heading(doc, "3. Option A — Oracle Cloud Always Free (Recommandée)", level=1, color=(31, 107, 58))

add_info_box(doc,
    "Oracle Cloud offre des VMs ARM permanentes et gratuites avec 24 Go de RAM — "
    "suffisant pour faire tourner Ollama complet sans aucune modification du code.",
    bg="E8F5E9", text_color=(31, 107, 58)
)

doc.add_paragraph()
add_heading(doc, "3.1 Ressources offertes gratuitement", level=2)
add_table(doc,
    ["Ressource", "Quota gratuit permanent", "Suffisant ?"],
    [
        ("CPU", "4 cœurs ARM Ampere A1", "Oui"),
        ("RAM", "24 Go", "Oui (mistral:7b = ~4-8 Go)"),
        ("Stockage", "200 Go (boot) + 2 volumes de 50 Go", "Oui"),
        ("Réseau", "10 To sortant/mois", "Oui"),
        ("IP publique", "2 IPs fixes gratuites", "Oui"),
    ],
    header_color="1F6B3A",
    alt_color="E8F5E9"
)

doc.add_paragraph()
add_heading(doc, "3.2 Architecture déployée", level=2)
doc.add_paragraph("Sur la VM Oracle Cloud, l'architecture locale est répliquée à l'identique :")

add_table(doc,
    ["Composant", "Service", "Port"],
    [
        ("Ollama (LLM + Embeddings)", "ollama serve", "11434 (interne)"),
        ("ChromaDB", "chromadb server", "8000 (interne)"),
        ("API RAG (FastAPI)", "uvicorn main:app", "8080 (exposé)"),
        ("Interface Web (optionnel)", "Streamlit / Gradio", "8501 (exposé)"),
    ]
)

doc.add_paragraph()
add_heading(doc, "3.3 Étapes de déploiement", level=2)
steps_a = [
    ("1", "Créer un compte Oracle Cloud", "cloud.oracle.com → Always Free → Instance ARM Ubuntu 22.04"),
    ("2", "Installer Ollama", "curl -fsSL https://ollama.com/install.sh | sh"),
    ("3", "Télécharger les modèles", "ollama pull nomic-embed-text && ollama pull mistral"),
    ("4", "Cloner le projet", "git clone <repo> && pip install -r requirements.txt"),
    ("5", "Créer l'API FastAPI", "Wrapper du pipeline existant (voir Section 5)"),
    ("6", "Configurer le pare-feu", "Ouvrir les ports 8080 et 8501 dans les Security Lists Oracle"),
    ("7", "Lancer en production", "systemd service pour Ollama + Gunicorn/Uvicorn"),
]
add_table(doc, ["Étape", "Action", "Détail"], steps_a)

doc.add_paragraph()
add_heading(doc, "3.4 Avantages / Inconvénients", level=2)
add_table(doc,
    ["Avantages", "Inconvénients"],
    [
        ("Données 100% locales sur le serveur, aucune fuite", "Nécessite une configuration initiale (SSH, pare-feu)"),
        ("Gratuit permanent, aucun plafond de requêtes", "Performances CPU limitées (ARM, pas de GPU)"),
        ("Aucune modification du code existant", "Temps de réponse ~5-15 sec par requête (CPU seul)"),
        ("Contrôle total sur les modèles utilisés", "Disponibilité dépend d'Oracle (SLA non garanti sur free tier)"),
    ],
    header_color="1F3864",
    alt_color="EBF3FB"
)

doc.add_paragraph()

# ════════════════════════════════════════
# OPTION B
# ════════════════════════════════════════
add_heading(doc, "4. Option B — Architecture Hybride (API LLM Gratuite)", level=1, color=(31, 56, 100))

add_info_box(doc,
    "Déployer uniquement l'application (FastAPI + ChromaDB) sur un cloud léger, "
    "et remplacer Ollama par une API LLM gratuite (Groq ou Gemini). "
    "Aucune RAM lourde requise — fonctionne sur n'importe quel free tier.",
    bg="D6E4F7", text_color=(31, 56, 100)
)

doc.add_paragraph()
add_heading(doc, "4.1 Schéma d'architecture", level=2)
add_table(doc,
    ["Couche", "Service Cloud", "Coût"],
    [
        ("Application + ChromaDB", "Render.com / Fly.io / Railway", "Gratuit (512 Mo RAM suffisant sans Ollama)"),
        ("Embeddings", "nomic-embed-text via Ollama local OU HuggingFace", "Gratuit"),
        ("Génération LLM", "Groq API (llama3-70b) OU Gemini Flash", "Gratuit (limites quotidiennes)"),
        ("Interface Web", "Streamlit Cloud / HuggingFace Spaces", "Gratuit"),
    ],
    header_color="1F3864",
    alt_color="EBF3FB"
)

doc.add_paragraph()
add_heading(doc, "4.2 Comparatif des API LLM gratuites", level=2)
add_table(doc,
    ["Provider", "Modèle", "Limite gratuite", "Qualité français", "Latence"],
    [
        ("Groq", "llama3-70b / mixtral-8x7b", "14 400 req/jour, 6 000 tokens/min", "Très bonne", "Très rapide (~1 sec)"),
        ("Google Gemini", "gemini-1.5-flash", "1 500 req/jour, 15 req/min", "Excellente", "Rapide (~2 sec)"),
        ("Cloudflare AI", "llama3, mistral", "10 000 req/jour", "Bonne", "Rapide (~2 sec)"),
        ("HuggingFace", "mistral-7b-instruct", "Appels limités, dépend du modèle", "Bonne", "Variable (3-10 sec)"),
    ],
    header_color="1A5276",
    alt_color="D6EAF8"
)

doc.add_paragraph()
add_heading(doc, "4.3 Modification du code pour Groq", level=2)
doc.add_paragraph("Une seule modification dans generation.py — remplacer _generate_ollama() :")
add_code_block(doc, "pip install groq")
add_code_block(doc,
    "from groq import Groq\n"
    "client = Groq(api_key=os.getenv('GROQ_API_KEY'))  # clé gratuite sur console.groq.com\n\n"
    "def _generate_groq(self, prompt: str) -> str:\n"
    "    response = client.chat.completions.create(\n"
    "        model='llama-3.3-70b-versatile',\n"
    "        messages=[{'role': 'user', 'content': prompt}],\n"
    "        temperature=0.1\n"
    "    )\n"
    "    return response.choices[0].message.content"
)

doc.add_paragraph()
add_heading(doc, "4.4 Avantages / Inconvénients", level=2)
add_table(doc,
    ["Avantages", "Inconvénients"],
    [
        ("Déploiement très simple et rapide", "Données de la requête envoyées à un tiers (Groq/Google)"),
        ("Scalabilité automatique", "Limite quotidienne de requêtes (14 400/jour pour Groq)"),
        ("Faible RAM requise (~256 Mo)", "Dépendance à la disponibilité du provider"),
        ("Génération très rapide (GPU cloud)", "Clé API nécessaire (même si gratuite)"),
    ],
    header_color="1F3864",
    alt_color="EBF3FB"
)

doc.add_paragraph()

# ════════════════════════════════════════
# OPTION C
# ════════════════════════════════════════
add_heading(doc, "5. Option C — VPS Payant (Production)", level=1, color=(100, 60, 0))

add_info_box(doc,
    "Pour une mise en production professionnelle avec haute disponibilité, "
    "un VPS dédié à 5-20 €/mois offre GPU/CPU suffisant pour Ollama avec SLA garanti.",
    bg="FFF8E1", text_color=(100, 60, 0)
)

doc.add_paragraph()
add_table(doc,
    ["Provider", "Offre", "RAM", "GPU", "Prix/mois", "Recommandé pour"],
    [
        ("Hetzner Cloud", "CPX31", "8 Go RAM", "Non", "~10 €", "Petite équipe, CPU suffisant"),
        ("OVHcloud", "VPS Value", "4 Go RAM", "Non", "~6 €", "Souveraineté française"),
        ("Scaleway", "GPU Dev1-S", "10 Go RAM", "NVIDIA P100", "~25 €", "Inférence rapide avec GPU"),
        ("Google Cloud", "e2-standard-2", "8 Go RAM", "Non", "~50 €", "Scalabilité, intégration GCP"),
        ("Azure", "B2s", "4 Go RAM", "Non", "~35 €", "Intégration Microsoft/Office 365"),
    ],
    header_color="7D6608",
    alt_color="FEF9E7"
)

doc.add_paragraph()

# ════════════════════════════════════════
# MODIFICATIONS CODE POUR CLOUD
# ════════════════════════════════════════
add_heading(doc, "6. Modifications du Code pour le Cloud", level=1, color=(31, 56, 100))

doc.add_paragraph(
    "Le code actuel est conçu en local. Pour un déploiement cloud, "
    "les adaptations suivantes sont nécessaires, quelle que soit l'option choisie."
)

doc.add_paragraph()
add_heading(doc, "6.1 Fichiers à modifier", level=2)
add_table(doc,
    ["Fichier", "Modification", "Option concernée"],
    [
        ("config.py", "Lire l'URL Ollama depuis variable d'env (OLLAMA_HOST)", "A, C"),
        ("config.py", "Remplacer client Ollama par client Groq/Gemini", "B"),
        ("generation.py", "Adapter _generate_ollama() pour appel API distant", "B"),
        ("indexing.py", "Adapter generate_embeddings() pour embeddings distants", "B"),
        ("main.py", "Créer une API FastAPI (endpoint /query)", "A, B, C"),
        ("requirements.txt", "Ajouter fastapi, uvicorn, groq (selon option)", "A, B, C"),
    ]
)

doc.add_paragraph()
add_heading(doc, "6.2 Exemple d'API FastAPI à créer", level=2)
add_code_block(doc,
    "# api.py — Wrapper FastAPI du pipeline existant\n"
    "from fastapi import FastAPI\n"
    "from pydantic import BaseModel\n"
    "from legal_rag.pipeline import LegalIngestionPipeline\n"
    "from legal_rag.generation import LegalAnswerGenerator\n\n"
    "app = FastAPI()\n"
    "pipeline = LegalIngestionPipeline()\n"
    "generator = LegalAnswerGenerator()\n\n"
    "class Query(BaseModel):\n"
    "    question: str\n"
    "    n_results: int = 3\n\n"
    "@app.post('/query')\n"
    "def query(q: Query):\n"
    "    results = pipeline.search(q.question, q.n_results)\n"
    "    answer = generator.generate_answer(q.question, results)\n"
    "    return {'answer': answer, 'sources': results['metadatas'][0]}"
)

doc.add_paragraph()

# ════════════════════════════════════════
# TABLEAU DE DECISION
# ════════════════════════════════════════
add_heading(doc, "7. Tableau de Décision", level=1, color=(31, 56, 100))

doc.add_paragraph(
    "Choisir l'option selon les priorités du projet :"
)
doc.add_paragraph()

add_table(doc,
    ["Critère", "Option A\n(Oracle Free)", "Option B\n(Hybride Groq)", "Option C\n(VPS Payant)"],
    [
        ("Coût mensuel", "0 €", "0 €", "5 – 25 €"),
        ("Confidentialité données", "Maximale (local)", "Partielle (LLM externe)", "Maximale (local)"),
        ("Temps de réponse", "5 – 15 sec", "1 – 3 sec", "1 – 5 sec"),
        ("Requêtes/jour", "Illimitées", "14 400 (Groq)", "Illimitées"),
        ("Modification du code", "Aucune", "Faible (2 fichiers)", "Aucune"),
        ("Complexité déploiement", "Moyenne (SSH/VM)", "Faible", "Faible"),
        ("RGPD / Souveraineté", "Complète", "Partielle", "Complète"),
        ("GPU disponible", "Non (ARM CPU)", "Oui (côté API)", "En option (~25 €)"),
    ],
    header_color="1F3864",
    alt_color="EBF3FB"
)

doc.add_paragraph()
add_info_box(doc,
    "Recommandation : Pour un projet d'entreprise avec contrainte de confidentialité "
    "→ Option A (Oracle Cloud Free). Pour un déploiement rapide sans contrainte de confidentialité "
    "→ Option B (Groq). Pour la production à grande échelle → Option C (Hetzner ou OVHcloud).",
    bg="E8F5E9", text_color=(31, 107, 58)
)

doc.add_paragraph()

# ════════════════════════════════════════
# PIED DE PAGE
# ════════════════════════════════════════
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = f"Fiche Technique — Déploiement Cloud RAG — {datetime.date.today().strftime('%d/%m/%Y')}"
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer_para.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

output_path = "/Users/seck/Documents/Corte/RAG_TP/tp-m2/fiche_deploiement_cloud.docx"
doc.save(output_path)
print(f"Document généré : {output_path}")
