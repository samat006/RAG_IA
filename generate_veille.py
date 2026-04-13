from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_table(doc, headers, rows, header_color="1F3864", alt_color="EBF3FB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = alt_color if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    return table

def add_badge(doc, text, bg="1F6B3A"):
    """Encadré coloré style badge."""
    p = doc.add_paragraph()
    run = p.add_run(f"  {text}  ")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    # Fond via shading de paragraphe
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg)
    pPr.append(shd)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ════════════════════════════════════════
# PAGE DE TITRE
# ════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("VEILLE TECHNOLOGIQUE")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 56, 100)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Justification des Technologies — Système RAG")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(70, 130, 180)

doc.add_paragraph()
line = doc.add_paragraph("─" * 60)
line.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

badge_p = doc.add_paragraph()
badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
badge_run = badge_p.add_run("  STACK 100% GRATUITE — OPEN SOURCE — LOCAL  ")
badge_run.font.bold = True
badge_run.font.size = Pt(12)
badge_run.font.color.rgb = RGBColor(255, 255, 255)
pPr = badge_p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), '1F6B3A')
pPr.append(shd)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = info.add_run(f"Projet d'Entreprise — Chatbot RAG Documentaire\nDate : {datetime.date.today().strftime('%d %B %Y')}")
run3.font.size = Pt(12)
run3.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ════════════════════════════════════════
# INTRODUCTION
# ════════════════════════════════════════
add_heading(doc, "1. Introduction", level=1, color=(31, 56, 100))

doc.add_paragraph(
    "Ce document présente la veille technologique justifiant les choix techniques "
    "effectués dans le cadre du développement d'un système de Retrieval-Augmented Generation (RAG). "
    "Ce système permet d'interroger en langage naturel n'importe quel corpus documentaire "
    "(PDF, XML, JSON) et de générer des réponses précises, sourcées et vérifiables."
)

note = doc.add_paragraph()
run_note = note.add_run(
    "Note : l'architecture est entièrement générique et s'adapte à tout type de document "
    "(juridique, RH, médical, technique, contractuel, etc.) en modifiant uniquement les "
    "patterns de détection de sections. L'ensemble de la stack est 100% gratuit, "
    "open source et s'exécute localement sans aucune dépendance à un service cloud payant."
)
run_note.font.italic = True
run_note.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# ════════════════════════════════════════
# ARCHITECTURE RAG
# ════════════════════════════════════════
add_heading(doc, "2. Architecture RAG — Paradigme retenu", level=1, color=(31, 56, 100))

doc.add_paragraph(
    "Le Retrieval-Augmented Generation (RAG) est le paradigme dominant en 2024-2025 pour "
    "les systèmes de question-réponse sur corpus documentaire. Il résout deux problèmes "
    "fondamentaux des LLM :"
)

for item in [
    "Hallucination : le LLM génère des informations fausses car il ne dispose pas du contexte exact.",
    "Coût de fine-tuning : adapter un modèle sur un corpus spécifique est coûteux et statique.",
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

doc.add_paragraph(
    "\nAu lieu de fine-tuner, le RAG récupère dynamiquement les passages les plus pertinents "
    "du corpus et les fournit comme contexte au LLM lors de la génération. "
    "Cette approche garantit des réponses ancrées dans les documents réels."
)

doc.add_paragraph()
add_heading(doc, "Flux de traitement", level=2)

add_table(doc, ["Étape", "Phase", "Description"], [
    ("1", "Ingestion", "PDF/XML/JSON → Extraction texte → Chunking → Embeddings Ollama → ChromaDB"),
    ("2", "Requête", "Question utilisateur → Embedding Ollama → Recherche Top-K voisins"),
    ("3", "Génération", "Chunks récupérés + Requête → Prompt LLM Ollama → Réponse citée"),
])

doc.add_paragraph()

# ════════════════════════════════════════
# OLLAMA
# ════════════════════════════════════════
add_heading(doc, "3. Ollama — Moteur LLM Local", level=1, color=(31, 56, 100))

add_heading(doc, "3.1 Présentation", level=2)
doc.add_paragraph(
    "Ollama est une plateforme open source qui permet de faire tourner des grands modèles de langage "
    "(LLM) en local, directement sur la machine, sans connexion internet ni clé API. "
    "Il joue le rôle de serveur d'inférence et expose une API REST compatible avec les standards du secteur. "
    "Dans ce projet, Ollama est utilisé pour deux tâches : la vectorisation des documents "
    "(modèle nomic-embed-text) et la génération de réponses (modèle mistral:7b)."
)

add_heading(doc, "3.2 Modèles utilisés", level=2)
add_table(doc,
    ["Modèle", "Tâche", "Taille", "Caractéristiques"],
    [
        ("nomic-embed-text", "Embeddings", "274 Mo", "768 dimensions, multilingue, optimisé pour la recherche sémantique."),
        ("mistral:7b", "Génération", "~4.1 Go", "7 milliards de paramètres, excellentes performances sur le français, instruction-tuned."),
    ]
)

add_heading(doc, "3.3 Justification du choix", level=2)
add_table(doc,
    ["Critère", "Détail"],
    [
        ("100% gratuit", "Aucun coût d'usage, aucun abonnement, aucune clé API. Coût = 0€ à l'infini."),
        ("Confidentialité totale", "Les données ne quittent jamais la machine locale. Conformité RGPD native, critique pour documents sensibles."),
        ("Zéro dépendance cloud", "Fonctionne sans internet. Aucun risque de coupure de service ou de changement tarifaire."),
        ("Qualité sur le français", "Mistral:7b est entraîné par Mistral AI (France) sur un large corpus francophone."),
        ("API standardisée", "Interface compatible OpenAI, facilitant la migration vers un autre provider si nécessaire."),
        ("Modèles swappables", "Changement de modèle en une ligne de config (ex: llama3, gemma2, qwen2.5)."),
    ]
)

add_heading(doc, "3.4 Alternatives écartées", level=2)
add_table(doc,
    ["Alternative", "Raison du rejet"],
    [
        ("Mistral AI (API cloud)", "Payant à l'usage — incompatible avec l'exigence 100% gratuit."),
        ("OpenAI GPT-4", "Payant, données hébergées hors EU, dépendance à un acteur tiers."),
        ("Groq (free tier)", "Gratuit mais limité (14 400 req/jour), données envoyées hors infrastructure."),
        ("HuggingFace Inference API", "Limite de requêtes sur le tier gratuit, moins performant sur le français."),
    ]
)

doc.add_paragraph()

# ════════════════════════════════════════
# CHROMADB
# ════════════════════════════════════════
add_heading(doc, "4. ChromaDB — Base Vectorielle", level=1, color=(31, 56, 100))

add_heading(doc, "4.1 Présentation", level=2)
doc.add_paragraph(
    "ChromaDB est une base de données vectorielle open source (Apache 2.0) conçue spécifiquement "
    "pour les applications LLM. Elle stocke les vecteurs (embeddings) et effectue des recherches "
    "par similarité sémantique (ANN) en quelques millisecondes, avec filtrage par métadonnées."
)

add_heading(doc, "4.2 Justification du choix", level=2)
add_table(doc,
    ["Critère", "Détail"],
    [
        ("100% gratuit", "Apache 2.0, open source, aucun coût de licence ni d'usage."),
        ("Zero-infrastructure", "PersistentClient = stockage sur fichiers locaux. Aucun serveur à déployer."),
        ("API Python native", "Intégration directe, gestion collections/filtres/requêtes très lisible."),
        ("Algorithme HNSW", "Recherche ANN en O(log n). Passage à l'échelle transparent."),
        ("Filtrage hybride", "Combinaison recherche vectorielle + filtres métadonnées (where={})."),
        ("Mode serveur disponible", "Migration locale → client/serveur sans changer le code applicatif."),
    ]
)

add_heading(doc, "4.3 Alternatives écartées", level=2)
add_table(doc,
    ["Alternative", "Raison du rejet"],
    [
        ("Pinecone", "SaaS uniquement, payant au-delà du free tier, données hors EU."),
        ("Weaviate", "Plus complexe à configurer, surdimensionné pour ce projet."),
        ("pgvector (PostgreSQL)", "Excellent en production, mais nécessite une instance PostgreSQL séparée."),
        ("FAISS (Meta)", "Très performant mais sans persistance native ni filtrage métadonnées."),
    ]
)

doc.add_paragraph()

# ════════════════════════════════════════
# PYMUPDF
# ════════════════════════════════════════
add_heading(doc, "5. PyMuPDF (fitz) — Extraction PDF", level=1, color=(31, 56, 100))

add_heading(doc, "5.1 Présentation", level=2)
doc.add_paragraph(
    "PyMuPDF est un binding Python de la librairie C++ MuPDF, leader de l'extraction PDF. "
    "Il donne accès aux coordonnées géométriques de chaque bloc de texte, ce qui permet "
    "la suppression automatique des en-têtes et pieds de page sans règles hardcodées."
)

add_heading(doc, "5.2 Justification du choix", level=2)
add_table(doc,
    ["Critère", "Détail"],
    [
        ("100% gratuit", "AGPL-3.0, open source, aucun coût."),
        ("Extraction géométrique", "Coordonnées (x0, y0, x1, y1) par bloc → suppression en-têtes/pieds par position (8-92% de la page)."),
        ("Robustesse", "Gère les vieux PDF, encodages exotiques, PDF scannés mieux que pdfminer ou pdfplumber."),
        ("Performance", "Cœur C++ (MuPDF) — 10× plus rapide que pdfplumber sur gros volumes."),
        ("Généricité", "Fonctionne sur tout PDF, quelle que soit sa structure ou son domaine."),
    ]
)

doc.add_paragraph()

# ════════════════════════════════════════
# STRATÉGIES DE CHUNKING
# ════════════════════════════════════════
add_heading(doc, "6. Stratégies de Chunking", level=1, color=(31, 56, 100))

doc.add_paragraph(
    "Le chunking est l'étape la plus critique d'un pipeline RAG. Un découpage naïf à taille fixe "
    "brise la cohérence sémantique et dégrade fortement la qualité de la recherche. "
    "Le projet implémente deux stratégies complémentaires, toutes deux 100% en code Python natif (gratuit)."
)

add_heading(doc, "6.1 Chunker Structurel + Récursif", level=2)
doc.add_paragraph(
    "Le StructuralChunker détecte les sections logiques via des expressions régulières configurables. "
    "Chaque section devient un chunk autonome. Si une section est trop longue, elle est redécoupée "
    "récursivement via une hiérarchie de séparateurs (paragraphe → ligne → phrase → virgule)."
)
add_table(doc,
    ["Avantage", "Explication"],
    [
        ("Cohérence sémantique", "Chaque chunk = unité logique du document, pas une fenêtre arbitraire."),
        ("Adaptabilité", "Patterns de section configurables pour tout type de document."),
        ("Overlap", "Chevauchement configurable pour éviter de couper des phrases clés."),
        ("Fallback robuste", "Si aucune structure détectée, le chunking récursif garantit un résultat correct."),
    ]
)

add_heading(doc, "6.2 Parent-Child Retriever", level=2)
doc.add_paragraph(
    "Cette stratégie avancée résout le trilemme du chunking : tension entre précision de la "
    "recherche vectorielle (petits chunks denses) et richesse du contexte fourni au LLM (grands blocs)."
)
add_table(doc,
    ["Composant", "Taille", "Rôle"],
    [
        ("Chunks enfants", "~400 caractères", "Indexés vectoriellement pour une recherche précise."),
        ("Chunks parents", "~2000 caractères", "Stockés comme contexte riche, récupérés via mapping."),
        ("Mapping", "parent_id dans métadonnées", "Chaque enfant pointe vers son parent."),
    ]
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    "Résultat : précision de recherche des petits vecteurs + richesse contextuelle des grands blocs, "
    "sans augmenter le coût d'indexation."
)
run.font.italic = True
run.font.color.rgb = RGBColor(31, 56, 100)

doc.add_paragraph()

# ════════════════════════════════════════
# SPACY
# ════════════════════════════════════════
add_heading(doc, "7. spaCy — Traitement NLP", level=1, color=(31, 56, 100))

doc.add_paragraph(
    "spaCy est la librairie NLP de référence en Python pour le traitement industriel du langage naturel. "
    "Elle est utilisée pour l'enrichissement automatique des métadonnées des documents ingérés."
)
add_table(doc,
    ["Critère", "Détail"],
    [
        ("100% gratuit", "MIT License, open source."),
        ("NER sur le français", "Modèle fr_core_news_lg — extraction automatique de dates, noms propres, organisations."),
        ("Production-ready", "Utilisé en production par des cabinets juridiques, banques, assurances, services publics."),
        ("Extensible", "Règles métier spécifiques via EntityRuler sans ré-entraînement."),
        ("Généricité", "Applicable à tout type de document pour enrichir les métadonnées automatiquement."),
    ]
)

doc.add_paragraph()

# ════════════════════════════════════════
# PYTHON-DOTENV
# ════════════════════════════════════════
add_heading(doc, "8. python-dotenv — Configuration", level=1, color=(31, 56, 100))

doc.add_paragraph(
    "python-dotenv gère les variables de configuration via un fichier .env non versionné. "
    "C'est le standard de facto pour la configuration des applications Python, "
    "conforme aux 12-Factor Apps. Dans ce projet, il gère les paramètres de configuration "
    "(chemins, noms de modèles, tailles de chunks) sans durcir ces valeurs dans le code."
)

doc.add_paragraph()

# ════════════════════════════════════════
# TABLEAU DE SYNTHESE
# ════════════════════════════════════════
add_heading(doc, "9. Tableau de Synthèse", level=1, color=(31, 56, 100))

add_table(doc,
    ["Technologie", "Modèle / Version", "Rôle", "Licence", "Coût"],
    [
        ("Python", "3.11+", "Langage principal", "PSF", "Gratuit"),
        ("Ollama", "mistral:7b", "Génération LLM", "MIT", "Gratuit"),
        ("Ollama", "nomic-embed-text", "Embeddings", "MIT / Apache 2.0", "Gratuit"),
        ("ChromaDB", "PersistentClient", "Base vectorielle", "Apache 2.0", "Gratuit"),
        ("PyMuPDF", "fitz", "Extraction PDF", "AGPL-3.0", "Gratuit"),
        ("spaCy", "fr_core_news_lg", "NLP / NER", "MIT", "Gratuit"),
        ("python-dotenv", "—", "Configuration", "BSD", "Gratuit"),
        ("RAG", "Structural + Parent-Child", "Paradigme Q&A", "—", "Gratuit"),
    ],
    header_color="1F6B3A",
    alt_color="E8F5E9"
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Coût total d'infrastructure : 0 € — aucune clé API, aucun abonnement cloud.")
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(31, 107, 58)

doc.add_paragraph()

# ════════════════════════════════════════
# PREREQUIS DEPLOIEMENT
# ════════════════════════════════════════
add_heading(doc, "10. Prérequis de Déploiement", level=1, color=(31, 56, 100))

doc.add_paragraph("Pour faire tourner le projet, les étapes suivantes sont nécessaires :")

steps = [
    ("1. Installer Ollama", "Télécharger depuis ollama.com — disponible Windows, macOS, Linux."),
    ("2. Télécharger les modèles", "ollama pull nomic-embed-text (274 Mo) + ollama pull mistral (~4.1 Go)"),
    ("3. Lancer Ollama", "ollama serve — tourne en arrière-plan sur localhost:11434"),
    ("4. Installer les dépendances", "pip install -r requirements.txt"),
    ("5. Lancer le projet", "python main.py --retriever recursive"),
]
add_table(doc, ["Étape", "Commande / Action"], steps)

doc.add_paragraph()
add_heading(doc, "Configuration matérielle minimale", level=2)
add_table(doc,
    ["Composant", "Minimum", "Recommandé"],
    [
        ("RAM", "8 Go", "16 Go"),
        ("Stockage", "10 Go libres", "20 Go libres"),
        ("CPU", "4 cœurs", "8 cœurs"),
        ("GPU", "Optionnel (CPU suffisant)", "GPU NVIDIA (inférence 5-10× plus rapide)"),
    ]
)

doc.add_paragraph()

# ════════════════════════════════════════
# GENERALISATION
# ════════════════════════════════════════
add_heading(doc, "11. Généricité du Système", level=1, color=(31, 56, 100))

doc.add_paragraph(
    "Le système est conçu pour être adapté à tout domaine documentaire. "
    "Le seul composant spécifique au domaine est le StructuralChunker (patterns de détection de sections). "
    "Tous les autres composants sont domain-agnostic."
)

add_table(doc,
    ["Domaine cible", "Modification requise", "Composants inchangés"],
    [
        ("Documents RH", "Patterns de section dans chunkers.py", "Ollama, ChromaDB, PyMuPDF, Pipeline"),
        ("Articles médicaux", "Patterns + prompt système dans generation.py", "Ollama, ChromaDB, PyMuPDF, Pipeline"),
        ("Documentation technique", "Patterns de section", "Tous les autres composants"),
        ("Contrats commerciaux", "Patterns + prompt système", "Tous les autres composants"),
        ("Appels d'offres (XML/JSON)", "Aucune modification", "Tous les composants"),
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    "Conclusion : cette stack 100% gratuite et open source offre les mêmes capacités "
    "qu'une solution cloud payante, avec en plus la confidentialité totale des données "
    "et zéro dépendance à des fournisseurs externes."
)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 56, 100)

doc.add_paragraph()

# ════════════════════════════════════════
# PIED DE PAGE
# ════════════════════════════════════════
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = f"Veille Technologique — Système RAG — Stack 100% Gratuite — {datetime.date.today().strftime('%d/%m/%Y')}"
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer_para.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

output_path = "/Users/seck/Documents/Corte/RAG_TP/tp-m2/veille_technologique.docx"
doc.save(output_path)
print(f"Document généré : {output_path}")
