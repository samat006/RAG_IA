"""
Génère un tableau d'évaluation RAG au format Word (.docx).
Usage : python generate_eval_table.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Configurations à tester ─────────────────────────────────────────────────
CONFIGS = [
    {"n_results": 4,  "max_distance": 0.6, "model": "mistral", "retriever": "recursive"},
    {"n_results": 4,  "max_distance": 0.9, "model": "mistral", "retriever": "recursive"},
    {"n_results": 8,  "max_distance": 0.6, "model": "mistral", "retriever": "recursive"},
    {"n_results": 8,  "max_distance": 0.9, "model": "mistral", "retriever": "recursive"},
    {"n_results": 8,  "max_distance": 1.2, "model": "mistral", "retriever": "recursive"},
    {"n_results": 12, "max_distance": 0.9, "model": "mistral", "retriever": "recursive"},
    {"n_results": 8,  "max_distance": 0.9, "model": "llama3",  "retriever": "recursive"},
    {"n_results": 8,  "max_distance": 0.9, "model": "mistral", "retriever": "parent-child"},
]

QUESTIONS = [
    "Quels hébergements sont disponibles en Centre-Corse ?",
    "Quelles activités outdoor propose la région ?",
    "Quels sont les monuments historiques à visiter ?",
]

HEADERS = [
    "#",
    "Question",
    "n_results",
    "max_distance",
    "Modèle",
    "Retriever",
    "Réponse (résumé)",
    "Sources trouvées",
    "Évaluation\n☐ Bon  ☐ Pas satisfait",
    "Notes",
]

COL_WIDTHS_CM = [0.8, 4.5, 1.3, 1.5, 1.5, 2.0, 5.5, 3.5, 3.2, 3.0]

BLUE  = "1E3A5F"
WHITE = "FFFFFF"
GREY  = "F0F2F5"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "D1D5DB")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_header_cell(cell, text):
    set_cell_bg(cell, BLUE)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_data_cell(cell, text, bold=False, center=False, grey=False):
    if grey:
        set_cell_bg(cell, GREY)
    set_cell_borders(cell)
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def main():
    doc = Document()

    # ── Marges réduites ──────────────────────────────────────────
    section = doc.sections[0]
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Titre ────────────────────────────────────────────────────
    title = doc.add_heading("Tableau d'évaluation — Système RAG", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    doc.add_paragraph("Master 2 · Travaux Pratiques RAG · Corpus Tourisme Centre-Corse")
    doc.add_paragraph()

    # ── Tableau ──────────────────────────────────────────────────
    num_rows = 1 + len(CONFIGS) * len(QUESTIONS)
    table = doc.add_table(rows=num_rows, cols=len(HEADERS))
    table.style = "Table Grid"

    # Largeurs
    for i, width in enumerate(COL_WIDTHS_CM):
        for row in table.rows:
            row.cells[i].width = Cm(width)

    # En-tête
    for j, h in enumerate(HEADERS):
        add_header_cell(table.rows[0].cells[j], h)

    # Lignes de données
    row_num = 1
    for ci, cfg in enumerate(CONFIGS):
        for qi, question in enumerate(QUESTIONS):
            row = table.rows[row_num]
            grey = (ci % 2 == 1)
            add_data_cell(row.cells[0], str(row_num), center=True, grey=grey)
            add_data_cell(row.cells[1], question, grey=grey)
            add_data_cell(row.cells[2], str(cfg["n_results"]),    center=True, grey=grey)
            add_data_cell(row.cells[3], str(cfg["max_distance"]), center=True, grey=grey)
            add_data_cell(row.cells[4], cfg["model"],             center=True, grey=grey)
            add_data_cell(row.cells[5], cfg["retriever"],         center=True, grey=grey)
            add_data_cell(row.cells[6], "", grey=grey)   # Réponse — à remplir
            add_data_cell(row.cells[7], "", grey=grey)   # Sources — à remplir
            add_data_cell(row.cells[8], "☐ Bon          ☐ Pas satisfait", center=True, grey=grey)
            add_data_cell(row.cells[9], "", grey=grey)   # Notes — à remplir
            row_num += 1

    # ── Légende ──────────────────────────────────────────────────
    doc.add_paragraph()
    legend = doc.add_paragraph()
    legend.add_run("Légende : ").bold = True
    legend.add_run(
        "n_results = nombre de chunks récupérés dans ChromaDB · "
        "max_distance = seuil L2 de pertinence (plus bas = plus strict) · "
        "Retriever : recursive = découpage standard, parent-child = hiérarchique"
    ).font.size = Pt(8)

    out = "evaluation_rag.docx"
    doc.save(out)
    print(f"✅  Fichier généré : {out}")


if __name__ == "__main__":
    main()
